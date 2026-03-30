#!/usr/bin/env python3
"""创建英语单词趣味练习题Word文档"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import random

# 单词列表
words = [
    "Dog", "Cat", "Mouse", "Lion", "Leopard",
    "Apple", "Cheese", "Bread", "Lollipop", "Milk",
    "Yogurt", "Elephant", "Donut"
]

# 单词对应的数据
word_data = {
    "Dog": {"emoji": "🐕", "color": "#FF8C00", "bg": "FFF8DC", "cn": "小狗"},
    "Cat": {"emoji": "🐱", "color": "#FFB6C1", "bg": "FFF0F5", "cn": "小猫"},
    "Mouse": {"emoji": "🐭", "color": "#A9A9A9", "bg": "F5F5F5", "cn": "老鼠"},
    "Lion": {"emoji": "🦁", "color": "#FFD700", "bg": "FFFACD", "cn": "狮子"},
    "Leopard": {"emoji": "🐆", "color": "#CD853F", "bg": "FAEBD7", "cn": "豹子"},
    "Apple": {"emoji": "🍎", "color": "#DC143C", "bg": "FFE4E1", "cn": "苹果"},
    "Cheese": {"emoji": "🧀", "color": "#FFD700", "bg": "FFFACD", "cn": "奶酪"},
    "Bread": {"emoji": "🍞", "color": "#DEB887", "bg": "FAEBD7", "cn": "面包"},
    "Lollipop": {"emoji": "🍭", "color": "#FF69B4", "bg": "FFE4E6", "cn": "棒棒糖"},
    "Milk": {"emoji": "🥛", "color": "#87CEEB", "bg": "E6F3FF", "cn": "牛奶", "use_image": True},
    "Yogurt": {"emoji": "🥄", "color": "#E91E63", "bg": "FCE4EC", "cn": "酸奶", "use_image": True},
    "Elephant": {"emoji": "🐘", "color": "#708090", "bg": "F5F5F5", "cn": "大象"},
    "Donut": {"emoji": "🍩", "color": "#FF69B4", "bg": "FFE4E6", "cn": "甜甜圈"}
}

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ['top', 'left', 'bottom', 'right']:
        if edge in kwargs:
            element = OxmlElement(f'w:{edge}')
            element.set(qn('w:val'), kwargs[edge].get('val', 'single'))
            element.set(qn('w:sz'), str(kwargs[edge].get('sz', 4)))
            element.set(qn('w:color'), kwargs[edge].get('color', '000000'))
            tcBorders.append(element)
    tcPr.append(tcBorders)

def create_practice_document():
    """创建趣味练习题文档"""
    doc = Document()
    
    # 设置页面
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)
    
    # ==================== 第一部分：连线题 ====================
    # 标题
    title1 = doc.add_paragraph()
    title1_run = title1.add_run("🎯 第一部分：连线题")
    title1_run.font.size = Pt(28)
    title1_run.font.bold = True
    title1_run.font.color.rgb = RGBColor(0x2E, 0x86, 0xC1)
    title1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title1.space_after = Pt(10)
    
    # 说明
    desc1 = doc.add_paragraph()
    desc1_run = desc1.add_run("请用线把英文单词和对应的图片连起来")
    desc1_run.font.size = Pt(14)
    desc1_run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
    desc1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    desc1.space_after = Pt(15)
    
    # 随机选择8个单词用于连线题
    random.seed(42)  # 固定种子，确保可重复
    line_words = random.sample(words, 8)
    
    # 创建连线题表格（左边单词，右边图片）
    line_table = doc.add_table(rows=8, cols=2)
    line_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    line_table.autofit = False
    
    img_dir = "images"
    
    for i, word in enumerate(line_words):
        # 左列：单词
        left_cell = line_table.cell(i, 0)
        left_cell.width = Inches(2.5)
        set_cell_shading(left_cell, "F0F8FF")
        set_cell_border(left_cell, 
                       top={'val': 'single', 'sz': 8, 'color': '87CEEB'},
                       bottom={'val': 'single', 'sz': 8, 'color': '87CEEB'},
                       left={'val': 'single', 'sz': 8, 'color': '87CEEB'},
                       right={'val': 'single', 'sz': 8, 'color': '87CEEB'})
        
        left_para = left_cell.paragraphs[0]
        left_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        left_run = left_para.add_run(f"{i+1}. {word}")
        left_run.font.size = Pt(24)
        left_run.font.bold = True
        left_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        
        # 右列：图片或emoji
        right_cell = line_table.cell(i, 1)
        right_cell.width = Inches(3.5)
        set_cell_shading(right_cell, "FFF8DC")
        set_cell_border(right_cell, 
                       top={'val': 'single', 'sz': 8, 'color': 'FFD700'},
                       bottom={'val': 'single', 'sz': 8, 'color': 'FFD700'},
                       left={'val': 'single', 'sz': 8, 'color': 'FFD700'},
                       right={'val': 'single', 'sz': 8, 'color': 'FFD700'})
        
        right_para = right_cell.paragraphs[0]
        right_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if word in ["Milk", "Yogurt"]:
            img_path = os.path.join(img_dir, f"{word.lower()}.png")
            if os.path.exists(img_path):
                run = right_para.add_run()
                run.add_picture(img_path, width=Inches(0.8))
            else:
                right_run = right_para.add_run(word_data[word]["emoji"])
                right_run.font.size = Pt(36)
        else:
            right_run = right_para.add_run(word_data[word]["emoji"])
            right_run.font.size = Pt(36)
    
    # 添加分页
    doc.add_page_break()
    
    # ==================== 第二部分：选择题 ====================
    title2 = doc.add_paragraph()
    title2_run = title2.add_run("📝 第二部分：选择题")
    title2_run.font.size = Pt(28)
    title2_run.font.bold = True
    title2_run.font.color.rgb = RGBColor(0x2E, 0x86, 0xC1)
    title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title2.space_after = Pt(10)
    
    # 说明
    desc2 = doc.add_paragraph()
    desc2_run = desc2.add_run("看图选择正确的英文单词，在○里画√")
    desc2_run.font.size = Pt(14)
    desc2_run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
    desc2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    desc2.space_after = Pt(15)
    
    # 创建选择题（5道题）
    random.seed(43)
    choice_questions = random.sample(words, 5)
    
    for i, correct_word in enumerate(choice_questions):
        # 创建题目表格
        q_table = doc.add_table(rows=1, cols=2)
        q_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        q_table.autofit = False
        
        # 左列：图片
        left_cell = q_table.cell(0, 0)
        left_cell.width = Inches(1.8)
        set_cell_shading(left_cell, word_data[correct_word]["bg"])
        set_cell_border(left_cell, 
                       top={'val': 'single', 'sz': 12, 'color': word_data[correct_word]["color"].replace('#', '')},
                       bottom={'val': 'single', 'sz': 12, 'color': word_data[correct_word]["color"].replace('#', '')},
                       left={'val': 'single', 'sz': 12, 'color': word_data[correct_word]["color"].replace('#', '')},
                       right={'val': 'single', 'sz': 12, 'color': word_data[correct_word]["color"].replace('#', '')})
        
        left_para = left_cell.paragraphs[0]
        left_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if correct_word in ["Milk", "Yogurt"]:
            img_path = os.path.join(img_dir, f"{correct_word.lower()}.png")
            if os.path.exists(img_path):
                run = left_para.add_run()
                run.add_picture(img_path, width=Inches(1.2))
            else:
                left_run = left_para.add_run(word_data[correct_word]["emoji"])
                left_run.font.size = Pt(60)
        else:
            left_run = left_para.add_run(word_data[correct_word]["emoji"])
            left_run.font.size = Pt(60)
        
        # 右列：选项
        right_cell = q_table.cell(0, 1)
        right_cell.width = Inches(5.2)
        set_cell_shading(right_cell, "FFFFFF")
        
        right_para = right_cell.paragraphs[0]
        right_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # 题号
        q_num_run = right_para.add_run(f"{i+1}. ")
        q_num_run.font.size = Pt(20)
        q_num_run.font.bold = True
        
        # 生成选项（正确答案 + 3个干扰项）
        all_words = [w for w in words if w != correct_word]
        wrong_options = random.sample(all_words, 3)
        options = [correct_word] + wrong_options
        random.shuffle(options)
        
        # 显示选项
        for j, option in enumerate(options):
            option_para = right_cell.add_paragraph()
            option_run = option_para.add_run(f"○ {option}")
            option_run.font.size = Pt(18)
            option_para.space_before = Pt(6)
        
        # 添加间距
        space = doc.add_paragraph()
        space.space_after = Pt(15)
    
    # 添加分页
    doc.add_page_break()
    
    # ==================== 第三部分：看图写单词 ====================
    title3 = doc.add_paragraph()
    title3_run = title3.add_run("✍️ 第三部分：看图写单词")
    title3_run.font.size = Pt(28)
    title3_run.font.bold = True
    title3_run.font.color.rgb = RGBColor(0x2E, 0x86, 0xC1)
    title3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title3.space_after = Pt(10)
    
    # 说明
    desc3 = doc.add_paragraph()
    desc3_run = desc3.add_run("看图片，在横线上写出正确的英文单词")
    desc3_run.font.size = Pt(14)
    desc3_run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
    desc3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    desc3.space_after = Pt(20)
    
    # 创建填空题表格（2列x4行）
    random.seed(44)
    fill_words = random.sample(words, 8)
    
    fill_table = doc.add_table(rows=4, cols=2)
    fill_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    fill_table.autofit = False
    
    for i, word in enumerate(fill_words):
        row_idx = i // 2
        col_idx = i % 2
        cell = fill_table.cell(row_idx, col_idx)
        cell.width = Inches(3.5)
        
        set_cell_shading(cell, word_data[word]["bg"])
        set_cell_border(cell, 
                       top={'val': 'single', 'sz': 10, 'color': word_data[word]["color"].replace('#', '')},
                       bottom={'val': 'single', 'sz': 10, 'color': word_data[word]["color"].replace('#', '')},
                       left={'val': 'single', 'sz': 10, 'color': word_data[word]["color"].replace('#', '')},
                       right={'val': 'single', 'sz': 10, 'color': word_data[word]["color"].replace('#', '')})
        
        # 清空默认段落
        cell.paragraphs[0].clear()
        
        # 添加题号和图片
        num_para = cell.paragraphs[0]
        num_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        num_run = num_para.add_run(f"{i+1}. ")
        num_run.font.size = Pt(16)
        num_run.font.bold = True
        
        if word in ["Milk", "Yogurt"]:
            img_path = os.path.join(img_dir, f"{word.lower()}.png")
            if os.path.exists(img_path):
                run = num_para.add_run()
                run.add_picture(img_path, width=Inches(0.9))
            else:
                emoji_run = num_para.add_run(word_data[word]["emoji"])
                emoji_run.font.size = Pt(48)
        else:
            emoji_run = num_para.add_run(word_data[word]["emoji"])
            emoji_run.font.size = Pt(48)
        
        # 添加填空线
        line_para = cell.add_paragraph()
        line_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        line_run = line_para.add_run("____________")
        line_run.font.size = Pt(20)
        line_run.font.color.rgb = RGBColor(0x95, 0xA5, 0xA6)
    
    # 添加分页
    doc.add_page_break()
    
    # ==================== 第四部分：分类游戏 ====================
    title4 = doc.add_paragraph()
    title4_run = title4.add_run("🎮 第四部分：分类游戏")
    title4_run.font.size = Pt(28)
    title4_run.font.bold = True
    title4_run.font.color.rgb = RGBColor(0x2E, 0x86, 0xC1)
    title4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title4.space_after = Pt(10)
    
    # 说明
    desc4 = doc.add_paragraph()
    desc4_run = desc4.add_run("请把下面的单词分类，写在对应的方框里")
    desc4_run.font.size = Pt(14)
    desc4_run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
    desc4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    desc4.space_after = Pt(20)
    
    # 分类表格
    category_table = doc.add_table(rows=2, cols=2)
    category_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    category_table.autofit = False
    
    categories = [
        ("🐾 动物", ["Dog", "Cat", "Mouse", "Lion", "Leopard", "Elephant"]),
        ("🍎 食物", ["Apple", "Cheese", "Bread", "Lollipop", "Milk", "Yogurt", "Donut"])
    ]
    
    for i, (category_name, category_words) in enumerate(categories):
        row_idx = i // 2
        col_idx = i % 2
        cell = category_table.cell(row_idx, col_idx)
        cell.width = Inches(3.8)
        
        colors = ["FFF0F5", "FFFACD"]
        set_cell_shading(cell, colors[i])
        set_cell_border(cell, 
                       top={'val': 'single', 'sz': 12, 'color': '000000'},
                       bottom={'val': 'single', 'sz': 12, 'color': '000000'},
                       left={'val': 'single', 'sz': 12, 'color': '000000'},
                       right={'val': 'single', 'sz': 12, 'color': '000000'})
        
        # 类别标题
        title_para = cell.paragraphs[0]
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(category_name)
        title_run.font.size = Pt(22)
        title_run.font.bold = True
        
        # 空白区域（让孩子填写）
        for _ in range(6):
            space_para = cell.add_paragraph()
            space_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            space_run = space_para.add_run("___________")
            space_run.font.size = Pt(14)
            space_run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    
    # 添加答案页
    doc.add_page_break()
    
    # 答案标题
    answer_title = doc.add_paragraph()
    answer_run = answer_title.add_run("📖 答案页")
    answer_run.font.size = Pt(28)
    answer_run.font.bold = True
    answer_run.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)
    answer_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    answer_title.space_after = Pt(20)
    
    # 连线题答案
    line_answer = doc.add_paragraph()
    line_answer_run = line_answer.add_run("第一部分：连线题答案")
    line_answer_run.font.size = Pt(18)
    line_answer_run.font.bold = True
    line_answer_run.font.color.rgb = RGBColor(0x2E, 0x86, 0xC1)
    
    answer_content = doc.add_paragraph()
    for i, word in enumerate(line_words):
        answer_content.add_run(f"{i+1}. {word} → {word_data[word]['cn']}\n")
        answer_content.runs[-1].font.size = Pt(14)
    
    doc.add_paragraph()
    
    # 选择题答案
    choice_answer = doc.add_paragraph()
    choice_answer_run = choice_answer.add_run("第二部分：选择题答案")
    choice_answer_run.font.size = Pt(18)
    choice_answer_run.font.bold = True
    choice_answer_run.font.color.rgb = RGBColor(0x2E, 0x86, 0xC1)
    
    choice_content = doc.add_paragraph()
    for i, word in enumerate(choice_questions):
        choice_content.add_run(f"{i+1}. {word}\n")
        choice_content.runs[-1].font.size = Pt(14)
    
    doc.add_paragraph()
    
    # 填空题答案
    fill_answer = doc.add_paragraph()
    fill_answer_run = fill_answer.add_run("第三部分：看图写单词答案")
    fill_answer_run.font.size = Pt(18)
    fill_answer_run.font.bold = True
    fill_answer_run.font.color.rgb = RGBColor(0x2E, 0x86, 0xC1)
    
    fill_content = doc.add_paragraph()
    for i, word in enumerate(fill_words):
        fill_content.add_run(f"{i+1}. {word}\n")
        fill_content.runs[-1].font.size = Pt(14)
    
    # 保存文档
    output_path = "英语单词趣味练习题.docx"
    doc.save(output_path)
    print(f"✅ 练习题文档已创建: {output_path}")
    print(f"📄 包含4个部分的趣味练习")
    print(f"🎯 连线题：8题")
    print(f"📝 选择题：5题")
    print(f"✍️ 填空题：8题")
    print(f"🎮 分类游戏：2个类别")
    return output_path

if __name__ == "__main__":
    create_practice_document()
