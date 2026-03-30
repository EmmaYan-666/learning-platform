#!/usr/bin/env python3
"""创建英语单词卡片Word文档"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 单词列表
words = [
    "Dog", "Cat", "Mouse", "Lion", "Leopard",
    "Apple", "Cheese", "Bread", "Lollipop", "Milk",
    "Yogurt", "Elephant", "Donut"
]

# 单词对应的emoji和颜色
word_data = {
    "Dog": {"emoji": "🐕", "color": "FFA500", "bg": "FFF5E6"},
    "Cat": {"emoji": "🐱", "color": "FFD700", "bg": "FFF9E6"},
    "Mouse": {"emoji": "🐭", "color": "808080", "bg": "F0F0F0"},
    "Lion": {"emoji": "🦁", "color": "FFD700", "bg": "FFF9E6"},
    "Leopard": {"emoji": "🐆", "color": "D2691E", "bg": "F5E6D3"},
    "Apple": {"emoji": "🍎", "color": "FF0000", "bg": "FFE6E6"},
    "Cheese": {"emoji": "🧀", "color": "FFD700", "bg": "FFF9E6"},
    "Bread": {"emoji": "🍞", "color": "DEB887", "bg": "F5EFE6"},
    "Lollipop": {"emoji": "🍭", "color": "FF69B4", "bg": "FFE6F2"},
    "Milk": {"emoji": "🥛", "color": "4169E1", "bg": "E6F0FF"},
    "Yogurt": {"emoji": "🥛", "color": "9370DB", "bg": "F5E6FF"},
    "Elephant": {"emoji": "🐘", "color": "808080", "bg": "F0F0F0"},
    "Donut": {"emoji": "🍩", "color": "FF69B4", "bg": "FFE6F2"}
}

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def create_card_document():
    """创建单词卡片文档"""
    doc = Document()
    
    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    # 添加标题
    title = doc.add_paragraph()
    title_run = title.add_run("🌈 英语单词卡片 🌈")
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.space_after = Pt(20)
    
    # 创建表格 (每行2个卡片)
    cards_per_row = 2
    rows_needed = (len(words) + cards_per_row - 1) // cards_per_row
    
    table = doc.add_table(rows=rows_needed, cols=cards_per_row)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 设置表格样式
    table.autofit = False
    
    # 填充卡片
    for i, word in enumerate(words):
        row_idx = i // cards_per_row
        col_idx = i % cards_per_row
        cell = table.cell(row_idx, col_idx)
        
        # 设置单元格大小
        cell.width = Inches(3.5)
        
        # 设置背景色
        bg_color = word_data.get(word, {}).get("bg", "FFFFFF")
        set_cell_shading(cell, bg_color)
        
        # 清空单元格默认段落
        cell.paragraphs[0].clear()
        
        # 添加emoji图片区域（大号emoji）
        emoji_para = cell.paragraphs[0]
        emoji_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        emoji_run = emoji_para.add_run(word_data.get(word, {}).get("emoji", "⭐"))
        emoji_run.font.size = Pt(72)
        
        # 添加单词
        word_para = cell.add_paragraph()
        word_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        word_run = word_para.add_run(word)
        word_run.font.size = Pt(36)
        word_run.font.bold = True
        # 设置字体颜色
        text_color = word_data.get(word, {}).get("color", "333333")
        word_run.font.color.rgb = RGBColor(int(text_color[0:2], 16), 
                                           int(text_color[2:4], 16), 
                                           int(text_color[4:6], 16))
        
        # 添加发音提示（音标）
        phonetic_para = cell.add_paragraph()
        phonetic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        phonetic_run = phonetic_para.add_run(f"/{word.lower()}/")
        phonetic_run.font.size = Pt(16)
        phonetic_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        
        # 添加空白段落作为间距
        space_para = cell.add_paragraph()
        space_para.add_run(" ")
    
    # 填充剩余空单元格
    total_cells = rows_needed * cards_per_row
    for i in range(len(words), total_cells):
        row_idx = i // cards_per_row
        col_idx = i % cards_per_row
        cell = table.cell(row_idx, col_idx)
        cell.width = Inches(3.5)
        set_cell_shading(cell, "FFFFFF")
    
    # 添加页脚说明
    footer = doc.add_paragraph()
    footer.space_before = Pt(20)
    footer_run = footer.add_run("💡 学习小贴士：每个单词大声朗读3遍，记住它们的样子和发音！")
    footer_run.font.size = Pt(12)
    footer_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 保存文档
    output_path = "英语单词卡片.docx"
    doc.save(output_path)
    print(f"✅ Word文档已创建: {output_path}")
    return output_path

if __name__ == "__main__":
    create_card_document()
