#!/usr/bin/env python3
"""
生成优化的短语和例句文档
- 使用表格布局，节省空间
- 内容更集中，减少页数
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
from word_list_manager import init_word_database, get_all_words, get_word_info

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ['top', 'left', 'bottom', 'right']:
        if edge in kwargs:
            element = OxmlElement(f'w:{edge}')
            element.set(qn('w:val'), kwargs[edge].get('val', 'single'))
            element.set(qn('w:sz'), str(kwargs[edge].get('sz', 4)))
            element.set(qn('w:color'), kwargs[edge].get('color', '000000'))
            tcBorders.append(element)
    tcPr.append(tcBorders)

def generate_phrases_document_optimized(words=None, output_filename="英语短语和例句_优化版.docx"):
    """生成优化的短语和例句文档（使用表格布局）"""
    db = init_word_database()
    
    if words is None:
        words = get_all_words()
    
    doc = Document()
    
    # 设置页面
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    
    # 标题
    title = doc.add_paragraph()
    title_run = title.add_run("📚 英语短语和例句学习")
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x2E, 0x86, 0xC1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.space_after = Pt(5)
    
    # 说明
    desc = doc.add_paragraph()
    desc_run = desc.add_run("💡 每天练习3-5个单词的短语和例句，轻松掌握英语！")
    desc_run.font.size = Pt(12)
    desc_run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    desc.space_after = Pt(10)
    
    # 按类别分组
    categories = {}
    for word in words:
        word_info = get_word_info(word)
        if word_info:
            cat = word_info["category"]
            if cat not in categories:
                categories[cat] = []
            categories[cat].append(word)
    
    # 类别名称映射
    category_names = {
        "animals": "🐾 动物篇",
        "food": "🍎 食物篇",
        "family": "👨‍👩‍👧 家庭篇",
        "colors": "🎨 颜色篇",
        "adjectives": "📏 形容词篇"
    }
    
    # 为每个类别生成内容
    for category, category_words in categories.items():
        # 类别标题
        cat_title = doc.add_paragraph()
        cat_run = cat_title.add_run(category_names.get(category, category.title()))
        cat_run.font.size = Pt(20)
        cat_run.font.bold = True
        cat_run.font.color.rgb = RGBColor(0xE7, 0x4C, 0x3C)
        cat_title.space_before = Pt(8)
        cat_title.space_after = Pt(6)
        
        # 为每个单词创建紧凑表格
        for word in category_words:
            word_info = get_word_info(word)
            
            # 创建单词表格（1行3列：单词信息 | 短语 | 例句）
            word_table = doc.add_table(rows=1, cols=3)
            word_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            word_table.autofit = False
            
            # 第1列：单词信息
            cell1 = word_table.cell(0, 0)
            cell1.width = Inches(1.5)
            set_cell_shading(cell1, word_info["bg"])
            set_cell_border(cell1,
                          top={'val': 'single', 'sz': 8, 'color': word_info["color"].replace('#', '')},
                          bottom={'val': 'single', 'sz': 8, 'color': word_info["color"].replace('#', '')},
                          left={'val': 'single', 'sz': 8, 'color': word_info["color"].replace('#', '')},
                          right={'val': 'single', 'sz': 8, 'color': word_info["color"].replace('#', '')})
            
            # Emoji和单词
            p1 = cell1.paragraphs[0]
            p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            emoji_run = p1.add_run(f"{word_info['emoji']}\n")
            emoji_run.font.size = Pt(48)  # 从24增加到48
            
            word_run = p1.add_run(f"{word}\n")
            word_run.font.size = Pt(20)  # 从16增加到20
            word_run.font.bold = True
            
            cn_run = p1.add_run(f"({word_info['chinese']})")
            cn_run.font.size = Pt(14)  # 从11增加到14
            cn_run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
            
            # 第2列：短语
            cell2 = word_table.cell(0, 1)
            cell2.width = Inches(2.8)
            set_cell_shading(cell2, "F8F9FA")
            set_cell_border(cell2,
                          top={'val': 'single', 'sz': 8, 'color': 'BDC3C7'},
                          bottom={'val': 'single', 'sz': 8, 'color': 'BDC3C7'},
                          left={'val': 'single', 'sz': 8, 'color': 'BDC3C7'},
                          right={'val': 'single', 'sz': 8, 'color': 'BDC3C7'})
            
            p2 = cell2.paragraphs[0]
            p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
            title_run = p2.add_run("📝 简单短语：\n")
            title_run.font.size = Pt(12)
            title_run.font.bold = True
            title_run.font.color.rgb = RGBColor(0x34, 0x98, 0xDB)
            
            for phrase in word_info["phrases"]:
                phrase_run = p2.add_run(f"• {phrase}\n")
                phrase_run.font.size = Pt(12)
            
            # 第3列：例句
            cell3 = word_table.cell(0, 2)
            cell3.width = Inches(3.7)
            set_cell_shading(cell3, "FFFFFF")
            set_cell_border(cell3,
                          top={'val': 'single', 'sz': 8, 'color': 'BDC3C7'},
                          bottom={'val': 'single', 'sz': 8, 'color': 'BDC3C7'},
                          left={'val': 'single', 'sz': 8, 'color': 'BDC3C7'},
                          right={'val': 'single', 'sz': 8, 'color': 'BDC3C7'})
            
            p3 = cell3.paragraphs[0]
            p3.alignment = WD_ALIGN_PARAGRAPH.LEFT
            title_run2 = p3.add_run("💬 实用例句：\n")
            title_run2.font.size = Pt(12)
            title_run2.font.bold = True
            title_run2.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)
            
            for sentence in word_info["sentences"]:
                sent_run = p3.add_run(f"• {sentence}\n")
                sent_run.font.size = Pt(12)
            
            # 添加小间距
            space = doc.add_paragraph()
            space.space_before = Pt(0)
            space.space_after = Pt(3)
        
        # 如果不是最后一个类别，添加分页
        if category != list(categories.keys())[-1]:
            doc.add_page_break()
    
    # 保存文档
    doc.save(output_filename)
    print(f"\n✅ 优化版短语和例句文档已生成: {output_filename}")
    print(f"📄 包含 {len(words)} 个单词的短语和例句")
    print(f"💡 使用表格布局，内容更紧凑，页数更少")
    return output_filename

if __name__ == "__main__":
    print("\n" + "="*60)
    print("📝 生成优化版短语和例句文档")
    print("="*60)
    
    generate_phrases_document_optimized()
    
    print("\n✅ 完成！")
