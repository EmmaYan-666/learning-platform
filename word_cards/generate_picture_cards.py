#!/usr/bin/env python3
"""
生成单词图片卡片
- 纯图片卡片，方便打印和剪裁
- 每张卡片仅包含图片和英文单词
- 统一尺寸，整齐排列
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
from word_list_manager import init_word_database, get_all_words, get_word_info

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ['top', 'left', 'bottom', 'right']:
        if edge in kwargs:
            element = OxmlElement(f'w:{edge}')
            element.set(qn('w:val'), kwargs[edge].get('val', 'single'))
            element.set(qn('w:sz'), str(kwargs[edge].get('sz', 4)))
            element.set(qn('w:color'), kwargs[edge].get('color', '000000'))
            tcBorders.append(element)
    tcPr.append(tcBorders)

def generate_picture_cards(words=None, output_filename="单词图片卡片.docx"):
    """生成单词图片卡片文档"""
    db = init_word_database()
    
    # 如果没有指定单词，使用所有单词
    if words is None:
        words = get_all_words()
    
    # 创建Word文档
    doc = Document()
    
    # 设置页面 - 横向
    section = doc.sections[0]
    section.page_width = Inches(11)  # Letter横向宽度
    section.page_height = Inches(8.5)  # Letter横向高度
    section.top_margin = Inches(0.3)
    section.bottom_margin = Inches(0.3)
    section.left_margin = Inches(0.3)
    section.right_margin = Inches(0.3)
    
    # 创建表格 (每行4个卡片)
    cards_per_row = 4
    rows_needed = (len(words) + cards_per_row - 1) // cards_per_row
    
    table = doc.add_table(rows=rows_needed, cols=cards_per_row)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    img_dir = "images"
    
    # 填充卡片
    for i, word in enumerate(words):
        word_info = get_word_info(word)
        if not word_info:
            print(f"⚠️  单词 '{word}' 不存在于数据库中")
            continue
        
        row_idx = i // cards_per_row
        col_idx = i % cards_per_row
        cell = table.cell(row_idx, col_idx)
        
        # 设置单元格宽度和高度（固定尺寸）
        cell.width = Inches(2.5)
        
        # 获取样式
        bg_color = word_info.get("bg", "FFFFFF")
        text_color = word_info.get("color", "#333333").replace('#', '')
        emoji = word_info.get("emoji", "⭐")
        
        # 设置背景色和边框
        set_cell_shading(cell, bg_color)
        set_cell_border(cell, 
                       top={'val': 'single', 'sz': 12, 'color': text_color},
                       bottom={'val': 'single', 'sz': 12, 'color': text_color},
                       left={'val': 'single', 'sz': 12, 'color': text_color},
                       right={'val': 'single', 'sz': 12, 'color': text_color})
        
        # 清空默认段落
        cell.paragraphs[0].clear()
        
        # 添加图片或emoji（调整尺寸）
        if word_info.get("use_image"):
            img_path = os.path.join(img_dir, f"{word.lower()}.png")
            if os.path.exists(img_path):
                img_para = cell.add_paragraph()
                img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = img_para.add_run()
                run.add_picture(img_path, width=Inches(1.5))
            else:
                emoji_para = cell.add_paragraph()
                emoji_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                emoji_run = emoji_para.add_run(emoji)
                emoji_run.font.size = Pt(80)
        else:
            emoji_para = cell.add_paragraph()
            emoji_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            emoji_run = emoji_para.add_run(emoji)
            emoji_run.font.size = Pt(80)
        
        # 添加英文单词
        word_para = cell.add_paragraph()
        word_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        word_run = word_para.add_run(word)
        word_run.font.size = Pt(32)
        word_run.font.bold = True
        word_run.font.color.rgb = RGBColor(int(text_color[0:2], 16), 
                                           int(text_color[2:4], 16), 
                                           int(text_color[4:6], 16))
    
    # 填充剩余空单元格
    total_cells = rows_needed * cards_per_row
    for i in range(len(words), total_cells):
        row_idx = i // cards_per_row
        col_idx = i % cards_per_row
        cell = table.cell(row_idx, col_idx)
        cell.width = Inches(2.5)
        set_cell_shading(cell, "FFFFFF")
    
    # 保存文档
    doc.save(output_filename)
    print(f"\n✅ 单词图片卡片已生成: {output_filename}")
    print(f"📄 包含 {len(words)} 张单词图片卡片")
    print(f"🖨️  横向打印，每页8张卡片（4行x2列）")
    print(f"✂️  沿彩色边框剪裁即可使用")
    return output_filename

if __name__ == "__main__":
    print("\n" + "="*60)
    print("🎨 生成单词图片卡片")
    print("="*60)
    
    generate_picture_cards()
    
    print("\n✅ 完成！")
