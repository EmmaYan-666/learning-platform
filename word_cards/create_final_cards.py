#!/usr/bin/env python3
"""创建英语单词卡片Word文档 - 更新版（酸奶和牛奶差异化）"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# 单词列表
words = [
    "Dog", "Cat", "Mouse", "Lion", "Leopard",
    "Apple", "Cheese", "Bread", "Lollipop", "Milk",
    "Yogurt", "Elephant", "Donut"
]

# 单词对应的数据（包含搜索关键词和颜色）
word_data = {
    "Dog": {"search": "cute+cartoon+dog+illustration", "color": "#FF8C00", "bg": "FFF8DC", "cn": "小狗"},
    "Cat": {"search": "cute+cartoon+cat+kitten", "color": "#FFB6C1", "bg": "FFF0F5", "cn": "小猫"},
    "Mouse": {"search": "cute+cartoon+mouse+illustration", "color": "#A9A9A9", "bg": "F5F5F5", "cn": "老鼠"},
    "Lion": {"search": "cute+cartoon+lion+illustration", "color": "#FFD700", "bg": "FFFACD", "cn": "狮子"},
    "Leopard": {"search": "cute+cartoon+leopard+illustration", "color": "#CD853F", "bg": "FAEBD7", "cn": "豹子"},
    "Apple": {"search": "cute+cartoon+red+apple+illustration", "color": "#DC143C", "bg": "FFE4E1", "cn": "苹果"},
    "Cheese": {"search": "cute+cartoon+cheese+illustration", "color": "#FFD700", "bg": "FFFACD", "cn": "奶酪"},
    "Bread": {"search": "cute+cartoon+bread+loaf+illustration", "color": "#DEB887", "bg": "FAEBD7", "cn": "面包"},
    "Lollipop": {"search": "cute+cartoon+lollipop+candy+illustration", "color": "#FF69B4", "bg": "FFE4E6", "cn": "棒棒糖"},
    "Milk": {"search": "cute+cartoon+milk+glass+illustration", "color": "#87CEEB", "bg": "E6F3FF", "cn": "牛奶"},  # 蓝色系
    "Yogurt": {"search": "cute+cartoon+yogurt+cup+illustration", "color": "#E91E63", "bg": "FCE4EC", "cn": "酸奶"},  # 粉红色系
    "Elephant": {"search": "cute+cartoon+elephant+illustration", "color": "#708090", "bg": "F5F5F5", "cn": "大象"},
    "Donut": {"search": "cute+cartoon+donut+illustration", "color": "#FF69B4", "bg": "FFE4E6", "cn": "甜甜圈"}
}

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ['top', 'left', 'bottom', 'right']:
        if edge in kwargs:
            element = OxmlElement(f'w:{edge}')
            element.set(qn('w:val'), kwargs[edge].get('val', 'single'))
            element.set(qn('w:sz'), str(kwargs[edge].get('sz', 4)))
            element.set(qn('w:color'), kwargs[edge].get('color', '000000'))
            tcBorders.append(element)
    tcPr.append(tcBorders)

def create_card_document_with_emoji():
    """创建使用emoji的单词卡片文档"""
    doc = Document()
    
    # 设置页面
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    
    # 添加标题
    title = doc.add_paragraph()
    title_run = title.add_run("🌟 英语单词学习卡片 🌟")
    title_run.font.size = Pt(36)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x2E, 0x86, 0xC1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.space_after = Pt(8)
    
    # 副标题
    subtitle = doc.add_paragraph()
    sub_run = subtitle.add_run("适合5岁儿童 | 共13个单词 | 打印后剪裁使用")
    sub_run.font.size = Pt(14)
    sub_run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.space_after = Pt(15)
    
    # 创建表格 (每行2个卡片)
    cards_per_row = 2
    rows_needed = (len(words) + cards_per_row - 1) // cards_per_row
    
    table = doc.add_table(rows=rows_needed, cols=cards_per_row)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    # Emoji映射（使用不同的emoji区分牛奶和酸奶）
    emoji_map = {
        "Dog": "🐕",
        "Cat": "🐱", 
        "Mouse": "🐭",
        "Lion": "🦁",
        "Leopard": "🐆",
        "Apple": "🍎",
        "Cheese": "🧀",
        "Bread": "🍞",
        "Lollipop": "🍭",
        "Milk": "🥛",      # 白色玻璃杯牛奶
        "Yogurt": "🫙",    # 使用罐子emoji代表酸奶
        "Elephant": "🐘",
        "Donut": "🍩"
    }
    
    # 填充卡片
    for i, word in enumerate(words):
        row_idx = i // cards_per_row
        col_idx = i % cards_per_row
        cell = table.cell(row_idx, col_idx)
        
        # 设置单元格宽度
        cell.width = Inches(3.75)
        
        # 获取单词数据
        data = word_data.get(word, {})
        bg_color = data.get("bg", "FFFFFF")
        text_color = data.get("color", "#333333").replace('#', '')
        emoji = emoji_map.get(word, "⭐")
        cn_word = data.get("cn", "")
        
        # 设置背景色和边框
        set_cell_shading(cell, bg_color)
        set_cell_border(cell, 
                       top={'val': 'single', 'sz': 16, 'color': text_color},
                       bottom={'val': 'single', 'sz': 16, 'color': text_color},
                       left={'val': 'single', 'sz': 16, 'color': text_color},
                       right={'val': 'single', 'sz': 16, 'color': text_color})
        
        # 清空默认段落
        cell.paragraphs[0].clear()
        
        # 添加顶部间距
        space_top = cell.paragraphs[0]
        space_top.add_run(" ")
        
        # 添加emoji（超大号）
        emoji_para = cell.add_paragraph()
        emoji_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        emoji_run = emoji_para.add_run(emoji)
        emoji_run.font.size = Pt(100)
        
        # 为酸奶和牛奶添加特殊说明
        if word == "Milk":
            note_para = cell.add_paragraph()
            note_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            note_run = note_para.add_run("(液体饮料)")
            note_run.font.size = Pt(12)
            note_run.font.color.rgb = RGBColor(0x87, 0xCE, 0xEB)
        elif word == "Yogurt":
            note_para = cell.add_paragraph()
            note_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            note_run = note_para.add_run("(半固体)")
            note_run.font.size = Pt(12)
            note_run.font.color.rgb = RGBColor(0xE9, 0x1E, 0x63)
        
        # 添加英文单词
        word_para = cell.add_paragraph()
        word_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        word_run = word_para.add_run(word)
        word_run.font.size = Pt(44)
        word_run.font.bold = True
        word_run.font.color.rgb = RGBColor(int(text_color[0:2], 16), 
                                           int(text_color[2:4], 16), 
                                           int(text_color[4:6], 16))
        
        # 添加中文翻译
        cn_para = cell.add_paragraph()
        cn_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cn_run = cn_para.add_run(cn_word)
        cn_run.font.size = Pt(26)
        cn_run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
        
        # 添加底部间距
        space_bottom = cell.add_paragraph()
        space_bottom.add_run(" ")
    
    # 填充剩余空单元格
    total_cells = rows_needed * cards_per_row
    for i in range(len(words), total_cells):
        row_idx = i // cards_per_row
        col_idx = i % cards_per_row
        cell = table.cell(row_idx, col_idx)
        cell.width = Inches(3.75)
        set_cell_shading(cell, "FFFFFF")
    
    # 添加分页
    doc.add_page_break()
    
    # 第二页：学习指导
    guide_title = doc.add_paragraph()
    guide_run = guide_title.add_run("📚 学习指导与建议")
    guide_run.font.size = Pt(28)
    guide_run.font.bold = True
    guide_run.font.color.rgb = RGBColor(0x2E, 0x86, 0xC1)
    guide_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    guide_title.space_after = Pt(20)
    
    # 学习方法
    methods = [
        ("🎯 看图说单词", "指着卡片上的图片，让孩子大声说出英文单词，反复练习"),
        ("👂 听音指图", "家长读出单词，让孩子找到对应的卡片"),
        ("🎲 卡片配对", "将所有卡片反面朝上，翻开两张，找出相同的单词"),
        ("✍️ 描写练习", "在卡片背面让孩子用铅笔描写单词"),
        ("🏠 生活应用", "在生活中找到这些物品，并说出对应的英文单词"),
        ("📖 故事创作", "用卡片上的物品编一个小故事")
    ]
    
    for title_text, description in methods:
        method_para = doc.add_paragraph()
        method_para.space_before = Pt(10)
        title_run = method_para.add_run(title_text)
        title_run.font.size = Pt(16)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0xE7, 0x4C, 0x3C)
        
        desc_para = doc.add_paragraph()
        desc_run = desc_para.add_run("     " + description)
        desc_run.font.size = Pt(14)
        desc_run.font.color.rgb = RGBColor(0x5D, 0x6D, 0x7E)
    
    # 特别说明：牛奶和酸奶的区别
    doc.add_paragraph()
    diff_title = doc.add_paragraph()
    diff_run = diff_title.add_run("💡 牛奶和酸奶的区别")
    diff_run.font.size = Pt(18)
    diff_run.font.bold = True
    diff_run.font.color.rgb = RGBColor(0x2E, 0x86, 0xC1)
    
    diff_items = [
        "Milk（牛奶）- 蓝色边框，白色液体，可以直接饮用",
        "Yogurt（酸奶）- 粉红色边框，半固体状态，需要用勺子吃"
    ]
    
    for item in diff_items:
        item_para = doc.add_paragraph()
        item_para.add_run("  ✓ " + item)
        item_para.runs[0].font.size = Pt(13)
        item_para.runs[0].font.color.rgb = RGBColor(0x5D, 0x6D, 0x7E)
    
    # 使用说明
    doc.add_paragraph()
    usage_title = doc.add_paragraph()
    usage_run = usage_title.add_run("📝 打印与使用说明")
    usage_run.font.size = Pt(20)
    usage_run.font.bold = True
    usage_run.font.color.rgb = RGBColor(0x2E, 0x86, 0xC1)
    
    usage_items = [
        "建议使用彩色打印机打印，效果更佳",
        "使用A4或Letter纸张打印",
        "打印后沿边框剪裁成独立卡片",
        "可以塑封保存，重复使用",
        "每天学习3-5个单词，循序渐进"
    ]
    
    for item in usage_items:
        item_para = doc.add_paragraph()
        item_para.add_run("  ✓ " + item)
        item_para.runs[0].font.size = Pt(13)
        item_para.runs[0].font.color.rgb = RGBColor(0x5D, 0x6D, 0x7E)
    
    # 保存文档
    output_path = "英语单词卡片_完整版.docx"
    doc.save(output_path)
    print(f"✅ Word文档已创建: {output_path}")
    print(f"📄 包含 {len(words)} 个单词卡片")
    print(f"📖 包含学习指导和打印说明")
    print(f"🎨 牛奶和酸奶已差异化处理")
    return output_path

if __name__ == "__main__":
    create_card_document_with_emoji()
