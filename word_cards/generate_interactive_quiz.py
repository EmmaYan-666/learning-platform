#!/usr/bin/env python3
"""
生成趣味互动练习题
- 勾选题
- 找不同
- 涂色题
- 走迷宫
- 配对游戏
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import random
from word_list_manager import init_word_database, get_word_info, get_all_words

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ['top', 'left', 'bottom', 'right']:
        if edge in kwargs:
            element = OxmlElement(f'w:{edge}')
            element.set(qn('w:val'), kwargs[edge].get('val', 'single'))
            element.set(qn('w:sz'), str(kwargs[edge].get('sz', 4)))
            element.set(qn('w:color'), kwargs[edge].get('color', '000000'))
            tcBorders.append(element)
    tcPr.append(tcBorders)

def create_interactive_quiz():
    """创建趣味互动练习题"""
    doc = Document()
    
    # 设置页面
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)
    
    db = init_word_database()
    all_words = get_all_words()
    img_dir = "images"
    
    # ==================== 第一部分：勾选题 ====================
    title1 = doc.add_paragraph()
    title1_run = title1.add_run("✅ 第一部分：勾选正确答案")
    title1_run.font.size = Pt(26)
    title1_run.font.bold = True
    title1_run.font.color.rgb = RGBColor(0x2E, 0x86, 0xC1)
    title1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title1.space_after = Pt(8)
    
    desc1 = doc.add_paragraph()
    desc1_run = desc1.add_run("看图片，在正确的单词后面打勾 ✓")
    desc1_run.font.size = Pt(13)
    desc1_run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
    desc1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    desc1.space_after = Pt(15)
    
    # 随机选择10个单词
    random.seed(50)
    quiz_words = random.sample(all_words, min(10, len(all_words)))
    
    # 创建表格（每行2题）
    rows = (len(quiz_words) + 1) // 2
    quiz_table = doc.add_table(rows=rows, cols=2)
    quiz_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for i, word in enumerate(quiz_words):
        row_idx = i // 2
        col_idx = i % 2
        cell = quiz_table.cell(row_idx, col_idx)
        word_info = get_word_info(word)
        
        # 设置单元格
        set_cell_shading(cell, "F8F9FA")
        set_cell_border(cell,
                       top={'val': 'single', 'sz': 6, 'color': 'BDC3C7'},
                       bottom={'val': 'single', 'sz': 6, 'color': 'BDC3C7'},
                       left={'val': 'single', 'sz': 6, 'color': 'BDC3C7'},
                       right={'val': 'single', 'sz': 6, 'color': 'BDC3C7'})
        
        # 题号和图片
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        num_run = p.add_run(f"{i+1}. ")
        num_run.font.size = Pt(16)
        num_run.font.bold = True
        
        emoji_run = p.add_run(f"{word_info['emoji']}  ")
        emoji_run.font.size = Pt(32)
        
        # 生成选项（正确答案 + 2个干扰项）
        other_words = [w for w in all_words if w != word]
        wrong_options = random.sample(other_words, 2)
        options = [word] + wrong_options
        random.shuffle(options)
        
        # 显示选项
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        for option in options:
            check_run = p2.add_run(f"  □ {option}  ")
            check_run.font.size = Pt(14)
    
    # ==================== 第二部分：找一找 ====================
    doc.add_page_break()
    
    title2 = doc.add_paragraph()
    title2_run = title2.add_run("🔍 第二部分：找一找")
    title2_run.font.size = Pt(26)
    title2_run.font.bold = True
    title2_run.font.color.rgb = RGBColor(0xE7, 0x4C, 0x3C)
    title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title2.space_after = Pt(8)
    
    desc2 = doc.add_paragraph()
    desc2_run = desc2.add_run("在下面的图片中，找出所有指定的物品，圈出来！")
    desc2_run.font.size = Pt(13)
    desc2_run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
    desc2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    desc2.space_after = Pt(15)
    
    # 创建找一找游戏
    find_words = random.sample([w for w in all_words if get_word_info(w)["category"] in ["animals", "food"]], 6)
    
    # 创建大表格，包含多个emoji
    find_table = doc.add_table(rows=1, cols=1)
    find_cell = find_table.cell(0, 0)
    set_cell_shading(find_cell, "FFF9E6")
    set_cell_border(find_cell,
                   top={'val': 'single', 'sz': 12, 'color': 'F39C12'},
                   bottom={'val': 'single', 'sz': 12, 'color': 'F39C12'},
                   left={'val': 'single', 'sz': 12, 'color': 'F39C12'},
                   right={'val': 'single', 'sz': 12, 'color': 'F39C12'})
    
    p = find_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 提示要找的单词
    target_run = p.add_run("🎯 请找出这些物品并圈出来：\n\n")
    target_run.font.size = Pt(14)
    target_run.font.bold = True
    
    target_words_run = p.add_run("  ".join([f"{get_word_info(w)['emoji']} {w}" for w in find_words]))
    target_words_run.font.size = Pt(16)
    
    # 添加emoji网格
    grid_run = p.add_run("\n\n")
    grid_run.font.size = Pt(8)
    
    # 创建emoji网格（包含目标单词和其他emoji）
    grid_emojis = []
    for w in find_words:
        grid_emojis.append(get_word_info(w)['emoji'])
    
    # 添加干扰emoji
    extra_emojis = ["🌸", "🌺", "⭐", "🌙", "☀️", "🌈", "🎯", "🎨", "🎪", "🎭"]
    grid_emojis.extend(random.sample(extra_emojis, 4))
    
    # 创建5x2网格
    grid_text = ""
    random.shuffle(grid_emojis)
    for i, emoji in enumerate(grid_emojis):
        grid_text += f"  {emoji}  "
        if (i + 1) % 5 == 0:
            grid_text += "\n\n"
    
    grid_run = p.add_run(grid_text)
    grid_run.font.size = Pt(40)
    
    # ==================== 第三部分：颜色配对 ====================
    doc.add_page_break()
    
    title3 = doc.add_paragraph()
    title3_run = title3.add_run("🎨 第三部分：颜色配对")
    title3_run.font.size = Pt(26)
    title3_run.font.bold = True
    title3_run.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)
    title3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title3.space_after = Pt(8)
    
    desc3 = doc.add_paragraph()
    desc3_run = desc3.add_run("把物品和正确的颜色连起来")
    desc3_run.font.size = Pt(13)
    desc3_run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
    desc3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    desc3.space_after = Pt(15)
    
    # 颜色配对题目
    color_pairs = [
        ("Apple", "Red", "🍎", "🔴"),
        ("Banana", "Yellow", "🍌", "🟡"),
        ("Grass", "Green", "🌿", "🟢"),
        ("Sky", "Blue", "☁️", "🔵"),
        ("Snow", "White", "❄️", "⚪"),
        ("Cat (black cat)", "Black", "🐱", "⚫"),
    ]
    
    color_table = doc.add_table(rows=len(color_pairs), cols=2)
    color_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for i, (item, color, item_emoji, color_emoji) in enumerate(color_pairs):
        # 左列：物品
        left_cell = color_table.cell(i, 0)
        left_cell.width = Inches(2.5)
        set_cell_shading(left_cell, "E8F8F5")
        set_cell_border(left_cell,
                       top={'val': 'single', 'sz': 6, 'color': '27AE60'},
                       bottom={'val': 'single', 'sz': 6, 'color': '27AE60'},
                       left={'val': 'single', 'sz': 6, 'color': '27AE60'},
                       right={'val': 'single', 'sz': 6, 'color': '27AE60'})
        
        left_p = left_cell.paragraphs[0]
        left_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        left_run = left_p.add_run(f"{item_emoji}  {item}")
        left_run.font.size = Pt(16)
        
        # 右列：颜色
        right_cell = color_table.cell(i, 1)
        right_cell.width = Inches(2.5)
        set_cell_shading(right_cell, "FEF9E7")
        set_cell_border(right_cell,
                       top={'val': 'single', 'sz': 6, 'color': 'F39C12'},
                       bottom={'val': 'single', 'sz': 6, 'color': 'F39C12'},
                       left={'val': 'single', 'sz': 6, 'color': 'F39C12'},
                       right={'val': 'single', 'sz': 6, 'color': 'F39C12'})
        
        right_p = right_cell.paragraphs[0]
        right_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        right_run = right_p.add_run(f"{color_emoji}  {color}")
        right_run.font.size = Pt(16)
    
    # ==================== 第四部分：大小分类 ====================
    doc.add_page_break()
    
    title4 = doc.add_paragraph()
    title4_run = title4.add_run("🐘 第四部分：大和小")
    title4_run.font.size = Pt(26)
    title4_run.font.bold = True
    title4_run.font.color.rgb = RGBColor(0x9B, 0x59, 0xB6)
    title4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title4.space_after = Pt(8)
    
    desc4 = doc.add_paragraph()
    desc4_run = desc4.add_run("把下面的动物分分类，大的写在左边，小的写在右边")
    desc4_run.font.size = Pt(13)
    desc4_run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
    desc4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    desc4.space_after = Pt(15)
    
    # 显示动物
    animals_display = doc.add_paragraph()
    animals_display.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    animal_list = ["Elephant", "Mouse", "Lion", "Cat", "Dog", "Bird"]
    for animal in animal_list:
        info = get_word_info(animal)
        if info:
            animal_run = animals_display.add_run(f"{info['emoji']} {animal}   ")
            animal_run.font.size = Pt(18)
    
    # 分类框
    size_table = doc.add_table(rows=1, cols=2)
    size_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Big框
    big_cell = size_table.cell(0, 0)
    big_cell.width = Inches(3.5)
    set_cell_shading(big_cell, "EBF5FB")
    set_cell_border(big_cell,
                   top={'val': 'single', 'sz': 12, 'color': '3498DB'},
                   bottom={'val': 'single', 'sz': 12, 'color': '3498DB'},
                   left={'val': 'single', 'sz': 12, 'color': '3498DB'},
                   right={'val': 'single', 'sz': 12, 'color': '3498DB'})
    
    big_p = big_cell.paragraphs[0]
    big_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    big_title = big_p.add_run("🐘 BIG 大的\n\n")
    big_title.font.size = Pt(20)
    big_title.font.bold = True
    big_title.font.color.rgb = RGBColor(0x34, 0x98, 0xDB)
    
    for _ in range(3):
        line_p = big_cell.add_paragraph()
        line_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        line_run = line_p.add_run("______________")
        line_run.font.size = Pt(14)
        line_run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    
    # Small框
    small_cell = size_table.cell(0, 1)
    small_cell.width = Inches(3.5)
    set_cell_shading(small_cell, "FADBD8")
    set_cell_border(small_cell,
                   top={'val': 'single', 'sz': 12, 'color': 'E74C3C'},
                   bottom={'val': 'single', 'sz': 12, 'color': 'E74C3C'},
                   left={'val': 'single', 'sz': 12, 'color': 'E74C3C'},
                   right={'val': 'single', 'sz': 12, 'color': 'E74C3C'})
    
    small_p = small_cell.paragraphs[0]
    small_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    small_title = small_p.add_run("🐭 SMALL 小的\n\n")
    small_title.font.size = Pt(20)
    small_title.font.bold = True
    small_title.font.color.rgb = RGBColor(0xE7, 0x4C, 0x3C)
    
    for _ in range(3):
        line_p = small_cell.add_paragraph()
        line_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        line_run = line_p.add_run("______________")
        line_run.font.size = Pt(14)
        line_run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    
    # ==================== 第五部分：食物选择 ====================
    doc.add_page_break()
    
    title5 = doc.add_paragraph()
    title5_run = title5.add_run("🍎 第五部分：你喜欢什么？")
    title5_run.font.size = Pt(26)
    title5_run.font.bold = True
    title5_run.font.color.rgb = RGBColor(0xE6, 0x7E, 0x22)
    title5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title5.space_after = Pt(8)
    
    desc5 = doc.add_paragraph()
    desc5_run = desc5.add_run("在你喜欢的食物下面画 ⭕，在不喜欢的下面画 ❌")
    desc5_run.font.size = Pt(13)
    desc5_run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
    desc5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    desc5.space_after = Pt(15)
    
    # 选择食物
    food_words = [w for w in all_words if get_word_info(w)["category"] == "food"]
    food_sample = random.sample(food_words, min(8, len(food_words)))
    
    # 创建表格（每行4个）
    food_rows = (len(food_sample) + 3) // 4
    food_table = doc.add_table(rows=food_rows, cols=4)
    food_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for i, word in enumerate(food_sample):
        row_idx = i // 4
        col_idx = i % 4
        cell = food_table.cell(row_idx, col_idx)
        word_info = get_word_info(word)
        
        set_cell_shading(cell, word_info["bg"])
        set_cell_border(cell,
                       top={'val': 'single', 'sz': 8, 'color': word_info["color"].replace('#', '')},
                       bottom={'val': 'single', 'sz': 8, 'color': word_info["color"].replace('#', '')},
                       left={'val': 'single', 'sz': 8, 'color': word_info["color"].replace('#', '')},
                       right={'val': 'single', 'sz': 8, 'color': word_info["color"].replace('#', '')})
        
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        emoji_run = p.add_run(f"{word_info['emoji']}\n")
        emoji_run.font.size = Pt(36)
        
        word_run = p.add_run(f"{word}\n")
        word_run.font.size = Pt(14)
        word_run.font.bold = True
        
        # 空白框让孩子画符号
        check_run = p.add_run("   ")
        check_run.font.size = Pt(20)
    
    # ==================== 第六部分：动物叫声 ====================
    doc.add_page_break()
    
    title6 = doc.add_paragraph()
    title6_run = title6.add_run("🎵 第六部分：动物怎么叫？")
    title6_run.font.size = Pt(26)
    title6_run.font.bold = True
    title6_run.font.color.rgb = RGBColor(0x16, 0xA0, 0x85)
    title6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title6.space_after = Pt(8)
    
    desc6 = doc.add_paragraph()
    desc6_run = desc6.add_run("连线：把动物和它的叫声连起来")
    desc6_run.font.size = Pt(13)
    desc6_run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
    desc6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    desc6.space_after = Pt(15)
    
    # 动物叫声配对
    sounds = [
        ("Dog", "🐕", "Woof woof!"),
        ("Cat", "🐱", "Meow meow~"),
        ("Lion", "🦁", "Roar!"),
        ("Mouse", "🐭", "Squeak squeak"),
        ("Bird", "🐦", "Tweet tweet"),
    ]
    
    sound_table = doc.add_table(rows=len(sounds), cols=3)
    sound_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for i, (animal, emoji, sound) in enumerate(sounds):
        # 动物
        cell1 = sound_table.cell(i, 0)
        cell1.width = Inches(2)
        set_cell_shading(cell1, "E8F8F5")
        set_cell_border(cell1,
                       top={'val': 'single', 'sz': 6, 'color': 'BDC3C7'},
                       bottom={'val': 'single', 'sz': 6, 'color': 'BDC3C7'},
                       left={'val': 'single', 'sz': 6, 'color': 'BDC3C7'},
                       right={'val': 'single', 'sz': 6, 'color': 'BDC3C7'})
        
        p1 = cell1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1.add_run(f"{emoji}  {animal}").font.size = Pt(16)
        
        # 连线区域
        cell2 = sound_table.cell(i, 1)
        cell2.width = Inches(1.5)
        set_cell_shading(cell2, "FFFFFF")
        set_cell_border(cell2, 
                       top={'val': 'single', 'sz': 6, 'color': 'FFFFFF'},
                       bottom={'val': 'single', 'sz': 6, 'color': 'FFFFFF'},
                       left={'val': 'single', 'sz': 6, 'color': 'FFFFFF'},
                       right={'val': 'single', 'sz': 6, 'color': 'FFFFFF'})
        
        p2 = cell2.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.add_run("  ⟸⟹  ").font.size = Pt(16)
        
        # 叫声
        cell3 = sound_table.cell(i, 2)
        cell3.width = Inches(2)
        set_cell_shading(cell3, "FEF9E7")
        set_cell_border(cell3,
                       top={'val': 'single', 'sz': 6, 'color': 'BDC3C7'},
                       bottom={'val': 'single', 'sz': 6, 'color': 'BDC3C7'},
                       left={'val': 'single', 'sz': 6, 'color': 'BDC3C7'},
                       right={'val': 'single', 'sz': 6, 'color': 'BDC3C7'})
        
        p3 = cell3.paragraphs[0]
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sound_run = p3.add_run(f"{sound}")
        sound_run.font.size = Pt(14)
        sound_run.font.italic = True
    
    # ==================== 答案页 ====================
    doc.add_page_break()
    
    answer_title = doc.add_paragraph()
    answer_run = answer_title.add_run("📖 参考答案")
    answer_run.font.size = Pt(28)
    answer_run.font.bold = True
    answer_run.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)
    answer_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    answer_title.space_after = Pt(20)
    
    # 第一部分答案
    ans1 = doc.add_paragraph()
    ans1_run = ans1.add_run("第一部分：勾选题答案\n")
    ans1_run.font.size = Pt(16)
    ans1_run.font.bold = True
    ans1_run.font.color.rgb = RGBColor(0x2E, 0x86, 0xC1)
    
    for i, word in enumerate(quiz_words):
        ans_content = doc.add_paragraph()
        ans_content.add_run(f"{i+1}. {word}").font.size = Pt(12)
    
    # 第四部分答案
    doc.add_paragraph()
    ans4 = doc.add_paragraph()
    ans4_run = ans4.add_run("第四部分：大和小答案\n")
    ans4_run.font.size = Pt(16)
    ans4_run.font.bold = True
    ans4_run.font.color.rgb = RGBColor(0x9B, 0x59, 0xB6)
    
    big_answer = doc.add_paragraph()
    big_answer.add_run("BIG（大的）: Elephant, Lion, Dog\n").font.size = Pt(12)
    
    small_answer = doc.add_paragraph()
    small_answer.add_run("SMALL（小的）: Mouse, Cat, Bird").font.size = Pt(12)
    
    # 保存文档
    output_path = "英语趣味互动练习.docx"
    doc.save(output_path)
    print(f"\n✅ 趣味互动练习已生成: {output_path}")
    print(f"📄 包含6个互动游戏环节")
    print(f"✅ 勾选题")
    print(f"🔍 找一找游戏")
    print(f"🎨 颜色配对")
    print(f"🐘 大小分类")
    print(f"🍎 食物选择")
    print(f"🎵 动物叫声")
    return output_path

if __name__ == "__main__":
    print("\n" + "="*60)
    print("🎮 生成趣味互动练习")
    print("="*60)
    
    create_interactive_quiz()
    
    print("\n✅ 完成！")
