#!/usr/bin/env python3
"""
S1 第一课：My Body - 身体部位单词
扩展词汇并生成学习材料
"""

from word_list_manager import add_word, print_word_list
from generate_picture_cards import generate_picture_cards
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def add_body_vocabulary():
    """添加身体部位词汇"""
    
    print("\n" + "="*60)
    print("📚 S1 第一课：My Body - 身体部位")
    print("="*60 + "\n")
    
    # 用户提供的词汇
    body_words = [
        ("Ear", "耳朵", "👂", "#F39C12", "FEF9E7", 1),
        ("Eye", "眼睛", "👁️", "#3498DB", "EBF5FB", 1),
        ("Head", "头", "🗣️", "#E74C3C", "FADBD8", 1),
        ("Hand", "手", "✋", "#F39C12", "FEF9E7", 1),
        ("Leg", "腿", "🦵", "#95A5A6", "F4F6F7", 1),
        ("Shoulder", "肩膀", "💪", "#E67E22", "FDF2E9", 2),
        ("Nose", "鼻子", "👃", "#E91E63", "FCE4EC", 1),
        ("Mouth", "嘴巴", "👄", "#E91E63", "FCE4EC", 1),
    ]
    
    # 扩展词汇（控制在15个以内）
    extended_words = [
        ("Arm", "手臂", "💪", "#E67E22", "FDF2E9", 1),
        ("Foot", "脚", "🦶", "#95A5A6", "F4F6F7", 1),
        ("Finger", "手指", "👆", "#F39C12", "FEF9E7", 1),
        ("Knee", "膝盖", "🦵", "#95A5A6", "F4F6F7", 1),
        ("Hair", "头发", "💇", "#8B4513", "FAEBD7", 1),
        ("Face", "脸", "😊", "#FFB6C1", "FFF0F5", 1),
        ("Neck", "脖子", "🗣️", "#E74C3C", "FADBD8", 1),
    ]
    
    all_words = body_words + extended_words
    
    print(f"📝 添加 {len(all_words)} 个身体部位单词...\n")
    
    for word, chinese, emoji, color, bg, difficulty in all_words:
        add_word(
            word=word,
            chinese=chinese,
            category="body",
            emoji=emoji,
            color=color,
            bg=bg,
            difficulty=difficulty,
            phrases=[
                f"My {word.lower()}.",
                f"Touch your {word.lower()}.",
                f"This is my {word.lower()}."
            ],
            sentences=[
                f"I have two {word.lower()}s." if word not in ["Hair", "Face"] else f"I have one {word.lower()}.",
                f"Point to your {word.lower()}.",
                f"Your {word.lower()} is beautiful."
            ]
        )
    
    print("\n✅ 身体部位词汇添加完成！\n")
    return [w[0] for w in all_words]

def create_body_matching_game(output_filename="S1_My_Body_连线题.docx"):
    """创建身体部位连线题"""
    doc = Document()
    
    # 设置页面
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    
    # 标题
    title = doc.add_paragraph()
    title_run = title.add_run("📚 S1 第一课：My Body")
    title_run.font.size = Pt(32)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x2E, 0x86, 0xC1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.space_after = Pt(8)
    
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run("🎯 连线题：身体部位")
    subtitle_run.font.size = Pt(24)
    subtitle_run.font.bold = True
    subtitle_run.font.color.rgb = RGBColor(0xE7, 0x4C, 0x3C)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.space_after = Pt(5)
    
    desc = doc.add_paragraph()
    desc_run = desc.add_run("请用线把英文单词和正确的图片连起来")
    desc_run.font.size = Pt(14)
    desc_run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    desc.space_after = Pt(20)
    
    # 身体部位数据
    body_parts = [
        ("Ear", "👂", "耳朵"),
        ("Eye", "👁️", "眼睛"),
        ("Head", "🗣️", "头"),
        ("Hand", "✋", "手"),
        ("Leg", "🦵", "腿"),
        ("Nose", "👃", "鼻子"),
        ("Mouth", "👄", "嘴巴"),
        ("Arm", "💪", "手臂"),
        ("Foot", "🦶", "脚"),
        ("Finger", "👆", "手指"),
    ]
    
    # 创建连线表格
    import random
    random.seed(100)
    
    # 打乱顺序
    left_items = body_parts.copy()
    right_items = body_parts.copy()
    random.shuffle(right_items)
    
    # 创建表格
    table = doc.add_table(rows=len(left_items), cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    def set_cell_shading(cell, color):
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), color)
        cell._tc.get_or_add_tcPr().append(shading)
    
    def set_cell_border(cell, **kwargs):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for edge in ['top', 'left', 'bottom', 'right']:
            if edge in kwargs:
                element = OxmlElement(f'w:{edge}')
                element.set(qn('w:val'), kwargs[edge].get('val', 'single'))
                element.set(qn('w:sz'), str(kwargs[edge].get('sz', 4)))
                element.set(qn('w:color'), kwargs[edge].get('color', '000000'))
                tcBorders.append(element)
        tcPr.append(tcBorders)
    
    for i, (left_word, right_word) in enumerate(zip(left_items, right_items)):
        # 左列：英文单词
        cell1 = table.cell(i, 0)
        cell1.width = Inches(2)
        set_cell_shading(cell1, "EBF5FB")
        set_cell_border(cell1,
                       top={'val': 'single', 'sz': 6, 'color': '3498DB'},
                       bottom={'val': 'single', 'sz': 6, 'color': '3498DB'},
                       left={'val': 'single', 'sz': 6, 'color': '3498DB'},
                       right={'val': 'single', 'sz': 6, 'color': '3498DB'})
        
        p1 = cell1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run1 = p1.add_run(f"{left_word[0]}")
        run1.font.size = Pt(18)
        run1.font.bold = True
        
        # 中列：连线区域
        cell2 = table.cell(i, 1)
        cell2.width = Inches(2)
        set_cell_shading(cell2, "FFFFFF")
        p2 = cell2.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run("  ⟸⟹  ")
        run2.font.size = Pt(18)
        
        # 右列：图片
        cell3 = table.cell(i, 2)
        cell3.width = Inches(2)
        set_cell_shading(cell3, "FFF9E6")
        set_cell_border(cell3,
                       top={'val': 'single', 'sz': 6, 'color': 'F39C12'},
                       bottom={'val': 'single', 'sz': 6, 'color': 'F39C12'},
                       left={'val': 'single', 'sz': 6, 'color': 'F39C12'},
                       right={'val': 'single', 'sz': 6, 'color': 'F39C12'})
        
        p3 = cell3.paragraphs[0]
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run3 = p3.add_run(f"{right_word[1]}  {right_word[2]}")
        run3.font.size = Pt(16)
    
    # 分页
    doc.add_page_break()
    
    # 第二部分：看图写单词
    title2 = doc.add_paragraph()
    title2_run = title2.add_run("✍️ 第二部分：看图写单词")
    title2_run.font.size = Pt(24)
    title2_run.font.bold = True
    title2_run.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)
    title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title2.space_after = Pt(5)
    
    desc2 = doc.add_paragraph()
    desc2_run = desc2.add_run("看图片，在横线上写出正确的英文单词")
    desc2_run.font.size = Pt(14)
    desc2_run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
    desc2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    desc2.space_after = Pt(20)
    
    # 创建填空表格
    fill_words = random.sample(body_parts, 8)
    
    fill_table = doc.add_table(rows=4, cols=2)
    fill_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for i, (word, emoji, chinese) in enumerate(fill_words):
        row_idx = i // 2
        col_idx = i % 2
        cell = fill_table.cell(row_idx, col_idx)
        cell.width = Inches(3.5)
        
        set_cell_shading(cell, "F0F8FF")
        set_cell_border(cell,
                       top={'val': 'single', 'sz': 8, 'color': '87CEEB'},
                       bottom={'val': 'single', 'sz': 8, 'color': '87CEEB'},
                       left={'val': 'single', 'sz': 8, 'color': '87CEEB'},
                       right={'val': 'single', 'sz': 8, 'color': '87CEEB'})
        
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        num_run = p.add_run(f"{i+1}. ")
        num_run.font.size = Pt(16)
        num_run.font.bold = True
        
        emoji_run = p.add_run(f"{emoji}  ")
        emoji_run.font.size = Pt(48)
        
        # 填空线
        line_para = cell.add_paragraph()
        line_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        line_run = line_para.add_run("______________")
        line_run.font.size = Pt(16)
        line_run.font.color.rgb = RGBColor(0x95, 0xA5, 0xA6)
    
    # 答案页
    doc.add_page_break()
    
    answer_title = doc.add_paragraph()
    answer_run = answer_title.add_run("📖 参考答案")
    answer_run.font.size = Pt(28)
    answer_run.font.bold = True
    answer_run.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)
    answer_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    answer_title.space_after = Pt(20)
    
    ans1 = doc.add_paragraph()
    ans1_run = ans1.add_run("连线题答案：\n")
    ans1_run.font.size = Pt(16)
    ans1_run.font.bold = True
    ans1_run.font.color.rgb = RGBColor(0x2E, 0x86, 0xC1)
    
    for i, (word, emoji, chinese) in enumerate(left_items):
        ans_content = doc.add_paragraph()
        ans_content.add_run(f"{i+1}. {word} → {chinese}").font.size = Pt(14)
    
    doc.add_paragraph()
    
    ans2 = doc.add_paragraph()
    ans2_run = ans2.add_run("填空题答案：\n")
    ans2_run.font.size = Pt(16)
    ans2_run.font.bold = True
    ans2_run.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)
    
    for i, (word, emoji, chinese) in enumerate(fill_words):
        ans_content = doc.add_paragraph()
        ans_content.add_run(f"{i+1}. {word}").font.size = Pt(14)
    
    # 保存文档
    doc.save(output_filename)
    print(f"\n✅ 连线题文档已生成: {output_filename}")
    print(f"📄 包含连线题和填空题")
    return output_filename

if __name__ == "__main__":
    # 添加身体部位词汇
    body_words = add_body_vocabulary()
    
    # 打印当前单词列表
    print_word_list()
    
    # 生成身体部位单词卡片
    print("\n" + "="*60)
    print("🎨 生成身体部位单词卡片")
    print("="*60)
    generate_picture_cards(words=body_words, output_filename="S1_My_Body_单词卡片.docx")
    
    # 生成连线题
    print("\n" + "="*60)
    print("🎯 生成连线题")
    print("="*60)
    create_body_matching_game()
    
    print("\n" + "="*60)
    print("✅ S1 第一课材料生成完成！")
    print("="*60 + "\n")
