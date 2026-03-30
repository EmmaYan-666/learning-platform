#!/usr/bin/env python3
"""创建精美的英语单词卡片Word文档（带SVG图片）"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
import io
import os

# 单词列表
words = [
    "Dog", "Cat", "Mouse", "Lion", "Leopard",
    "Apple", "Cheese", "Bread", "Lollipop", "Milk",
    "Yogurt", "Elephant", "Donut"
]

# 单词对应的数据
word_data = {
    "Dog": {"emoji": "🐕", "color": "#FF8C00", "bg": "#FFF8DC", "cn": "小狗"},
    "Cat": {"emoji": "🐱", "color": "#FFB6C1", "bg": "#FFF0F5", "cn": "小猫"},
    "Mouse": {"emoji": "🐭", "color": "#A9A9A9", "bg": "#F5F5F5", "cn": "老鼠"},
    "Lion": {"emoji": "🦁", "color": "#FFD700", "bg": "#FFFACD", "cn": "狮子"},
    "Leopard": {"emoji": "🐆", "color": "#CD853F", "bg": "#FAEBD7", "cn": "豹子"},
    "Apple": {"emoji": "🍎", "color": "#DC143C", "bg": "#FFE4E1", "cn": "苹果"},
    "Cheese": {"emoji": "🧀", "color": "#FFD700", "bg": "#FFFACD", "cn": "奶酪"},
    "Bread": {"emoji": "🍞", "color": "#DEB887", "bg": "#FAEBD7", "cn": "面包"},
    "Lollipop": {"emoji": "🍭", "color": "#FF69B4", "bg": "#FFE4E6", "cn": "棒棒糖"},
    "Milk": {"emoji": "🥛", "color": "#87CEEB", "bg": "#F0F8FF", "cn": "牛奶"},
    "Yogurt": {"emoji": "🥛", "color": "#DDA0DD", "bg": "#FFF0FF", "cn": "酸奶"},
    "Elephant": {"emoji": "🐘", "color": "#708090", "bg": "#F5F5F5", "cn": "大象"},
    "Donut": {"emoji": "🍩", "color": "#FF69B4", "bg": "#FFE4E6", "cn": "甜甜圈"}
}

def create_svg_emoji(emoji, size=200):
    """创建包含emoji的SVG图片"""
    svg_template = f'''<?xml version="1.0" encoding="UTF-8"?>
<svg xmlns="http://www.w3.org/2000/svg" width="{size}" height="{size}">
  <rect width="100%" height="100%" fill="white" fill-opacity="0"/>
  <text x="50%" y="50%" 
        font-size="{size * 0.7}" 
        text-anchor="middle" 
        dominant-baseline="central"
        font-family="Apple Color Emoji, Segoe UI Emoji, Noto Color Emoji, sans-serif">
    {emoji}
  </text>
</svg>'''
    return svg_template

def svg_to_png(svg_content, output_path, size=400):
    """将SVG转换为PNG"""
    try:
        png_data = cairosvg.svg2png(bytestring=svg_content.encode('utf-8'), 
                                      output_width=size, 
                                      output_height=size)
        with open(output_path, 'wb') as f:
            f.write(png_data)
        return True
    except Exception as e:
        print(f"SVG转换失败: {e}")
        return False

def create_simple_image(emoji, output_path, size=400):
    """创建简单的图片（不依赖cairosvg）"""
    try:
        # 创建白色背景图片
        img = Image.new('RGB', (size, size), 'white')
        img.save(output_path, 'PNG')
        print(f"创建简单图片: {output_path}")
        return True
    except Exception as e:
        print(f"创建图片失败: {e}")
        return False

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ['top', 'left', 'bottom', 'right']:
        if edge in kwargs:
            element = OxmlElement(f'w:{edge}')
            element.set(qn('w:val'), kwargs[edge].get('val', 'single'))
            element.set(qn('w:sz'), str(kwargs[edge].get('sz', 4)))
            element.set(qn('w:color'), kwargs[edge].get('color', '000000'))
            tcBorders.append(element)
    tcPr.append(tcBorders)

def create_card_document():
    """创建单词卡片文档"""
    doc = Document()
    
    # 设置页面
    section = doc.sections[0]
    section.page_width = Inches(8.5)  # Letter宽度
    section.page_height = Inches(11)  # Letter高度
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)
    
    # 添加标题
    title = doc.add_paragraph()
    title_run = title.add_run("🌟 英语单词学习卡片 🌟")
    title_run.font.size = Pt(32)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x2E, 0x86, 0xC1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.space_after = Pt(15)
    
    # 副标题
    subtitle = doc.add_paragraph()
    sub_run = subtitle.add_run("适合5岁儿童 | 共13个单词")
    sub_run.font.size = Pt(14)
    sub_run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.space_after = Pt(20)
    
    # 创建表格 (每行2个卡片)
    cards_per_row = 2
    rows_needed = (len(words) + cards_per_row - 1) // cards_per_row
    
    table = doc.add_table(rows=rows_needed, cols=cards_per_row)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    # 创建图片目录
    img_dir = "images"
    os.makedirs(img_dir, exist_ok=True)
    
    # 填充卡片
    for i, word in enumerate(words):
        row_idx = i // cards_per_row
        col_idx = i % cards_per_row
        cell = table.cell(row_idx, col_idx)
        
        # 设置单元格宽度
        cell.width = Inches(3.6)
        
        # 获取单词数据
        data = word_data.get(word, {})
        bg_color = data.get("bg", "#FFFFFF").replace('#', '')
        text_color = data.get("color", "#333333").replace('#', '')
        emoji = data.get("emoji", "⭐")
        cn_word = data.get("cn", "")
        
        # 设置背景色
        set_cell_shading(cell, bg_color)
        
        # 设置边框
        set_cell_border(cell, 
                       top={'val': 'single', 'sz': 12, 'color': text_color},
                       bottom={'val': 'single', 'sz': 12, 'color': text_color},
                       left={'val': 'single', 'sz': 12, 'color': text_color},
                       right={'val': 'single', 'sz': 12, 'color': text_color})
        
        # 清空默认段落
        cell.paragraphs[0].clear()
        
        # 添加顶部间距
        space_top = cell.paragraphs[0]
        space_top.add_run(" ")
        
        # 添加emoji（大号）
        emoji_para = cell.add_paragraph()
        emoji_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        emoji_run = emoji_para.add_run(emoji)
        emoji_run.font.size = Pt(80)
        
        # 添加英文单词
        word_para = cell.add_paragraph()
        word_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        word_run = word_para.add_run(word)
        word_run.font.size = Pt(40)
        word_run.font.bold = True
        word_run.font.color.rgb = RGBColor(int(text_color[0:2], 16), 
                                           int(text_color[2:4], 16), 
                                           int(text_color[4:6], 16))
        
        # 添加中文翻译
        cn_para = cell.add_paragraph()
        cn_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cn_run = cn_para.add_run(cn_word)
        cn_run.font.size = Pt(24)
        cn_run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
        
        # 添加底部间距
        space_bottom = cell.add_paragraph()
        space_bottom.add_run(" ")
    
    # 填充剩余空单元格
    total_cells = rows_needed * cards_per_row
    for i in range(len(words), total_cells):
        row_idx = i // cards_per_row
        col_idx = i % cards_per_row
        cell = table.cell(row_idx, col_idx)
        cell.width = Inches(3.6)
        set_cell_shading(cell, "FFFFFF")
    
    # 添加学习建议
    doc.add_paragraph()
    tips = doc.add_paragraph()
    tips.space_before = Pt(15)
    tips_title = tips.add_run("📚 学习建议：")
    tips_title.font.size = Pt(14)
    tips_title.font.bold = True
    tips_title.font.color.rgb = RGBColor(0x2E, 0x86, 0xC1)
    
    tip_items = [
        "看图说单词：指着图片，大声说出英文单词",
        "跟读练习：跟着爸爸妈妈读3遍",
        "找一找：在生活中找到这些物品",
        "卡片游戏：把卡片翻过来，猜猜是什么"
    ]
    
    for tip in tip_items:
        tip_para = doc.add_paragraph()
        tip_para.add_run("  ✦ " + tip)
        tip_para.runs[0].font.size = Pt(12)
        tip_para.runs[0].font.color.rgb = RGBColor(0x5D, 0x6D, 0x7E)
    
    # 保存文档
    output_path = "英语单词卡片.docx"
    doc.save(output_path)
    print(f"✅ Word文档已创建: {output_path}")
    print(f"📄 包含 {len(words)} 个单词卡片")
    return output_path

if __name__ == "__main__":
    create_card_document()
