#!/usr/bin/env python3
"""
生成英语单词卡片
- 自动从数据库读取单词
- 生成Word文档
- 包含短语和例句
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
from word_list_manager import init_word_database, get_all_words, get_word_info

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ['top', 'left', 'bottom', 'right']:
        if edge in kwargs:
            element = OxmlElement(f'w:{edge}')
            element.set(qn('w:val'), kwargs[edge].get('val', 'single'))
            element.set(qn('w:sz'), str(kwargs[edge].get('sz', 4)))
            element.set(qn('w:color'), kwargs[edge].get('color', '000000'))
            tcBorders.append(element)
    tcPr.append(tcBorders)

def generate_word_cards(words=None, output_filename="英语单词卡片_自定义.docx"):
    """生成单词卡片Word文档"""
    db = init_word_database()
    
    # 如果没有指定单词，使用所有单词
    if words is None:
        words = get_all_words()
    
    # 创建Word文档
    doc = Document()
    
    # 设置页面
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    
    # 添加标题
    title = doc.add_paragraph()
    title_run = title.add_run("🌟 英语单词学习卡片 🌟")
    title_run.font.size = Pt(36)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x2E, 0x86, 0xC1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.space_after = Pt(8)
    
    # 副标题
    subtitle = doc.add_paragraph()
    sub_run = subtitle.add_run(f"适合5岁儿童 | 共{len(words)}个单词 | 打印后剪裁使用")
    sub_run.font.size = Pt(14)
    sub_run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.space_after = Pt(15)
    
    # 创建表格 (每行2个卡片)
    cards_per_row = 2
    rows_needed = (len(words) + cards_per_row - 1) // cards_per_row
    
    table = doc.add_table(rows=rows_needed, cols=cards_per_row)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    img_dir = "images"
    
    # 填充卡片
    for i, word in enumerate(words):
        word_info = get_word_info(word)
        if not word_info:
            print(f"⚠️  单词 '{word}' 不存在于数据库中")
            continue
        
        row_idx = i // cards_per_row
        col_idx = i % cards_per_row
        cell = table.cell(row_idx, col_idx)
        
        # 设置单元格宽度
        cell.width = Inches(3.75)
        
        # 获取样式
        bg_color = word_info.get("bg", "FFFFFF")
        text_color = word_info.get("color", "#333333").replace('#', '')
        chinese = word_info.get("chinese", "")
        emoji = word_info.get("emoji", "⭐")
        
        # 设置背景色和边框
        set_cell_shading(cell, bg_color)
        set_cell_border(cell, 
                       top={'val': 'single', 'sz': 16, 'color': text_color},
                       bottom={'val': 'single', 'sz': 16, 'color': text_color},
                       left={'val': 'single', 'sz': 16, 'color': text_color},
                       right={'val': 'single', 'sz': 16, 'color': text_color})
        
        # 清空默认段落
        cell.paragraphs[0].clear()
        
        # 添加顶部间距
        space_top = cell.paragraphs[0]
        space_top.add_run(" ")
        
        # 添加图片或emoji
        if word_info.get("use_image"):
            img_path = os.path.join(img_dir, f"{word.lower()}.png")
            if os.path.exists(img_path):
                img_para = cell.add_paragraph()
                img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = img_para.add_run()
                run.add_picture(img_path, width=Inches(1.8))
            else:
                emoji_para = cell.add_paragraph()
                emoji_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                emoji_run = emoji_para.add_run(emoji)
                emoji_run.font.size = Pt(100)
        else:
            emoji_para = cell.add_paragraph()
            emoji_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            emoji_run = emoji_para.add_run(emoji)
            emoji_run.font.size = Pt(100)
        
        # 添加英文单词
        word_para = cell.add_paragraph()
        word_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        word_run = word_para.add_run(word)
        word_run.font.size = Pt(48)  # 增大字号
        word_run.font.bold = True
        word_run.font.color.rgb = RGBColor(int(text_color[0:2], 16), 
                                           int(text_color[2:4], 16), 
                                           int(text_color[4:6], 16))
        
        # 添加底部间距
        space_bottom = cell.add_paragraph()
        space_bottom.add_run(" ")
    
    # 填充剩余空单元格
    total_cells = rows_needed * cards_per_row
    for i in range(len(words), total_cells):
        row_idx = i // cards_per_row
        col_idx = i % cards_per_row
        cell = table.cell(row_idx, col_idx)
        cell.width = Inches(3.75)
        set_cell_shading(cell, "FFFFFF")
    
    # 保存文档
    doc.save(output_filename)
    print(f"\n✅ 单词卡片已生成: {output_filename}")
    print(f"📄 包含 {len(words)} 个单词卡片")
    return output_filename

def generate_phrases_document(words=None, output_filename="英语短语和例句.docx"):
    """生成短语和例句文档"""
    db = init_word_database()
    
    if words is None:
        words = get_all_words()
    
    doc = Document()
    
    # 设置页面
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    
    # 标题
    title = doc.add_paragraph()
    title_run = title.add_run("📚 英语短语和例句学习")
    title_run.font.size = Pt(32)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x2E, 0x86, 0xC1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.space_after = Pt(20)
    
    # 说明
    desc = doc.add_paragraph()
    desc_run = desc.add_run("💡 这些短语和例句适合家长和孩子一起练习，帮助巩固单词记忆！")
    desc_run.font.size = Pt(14)
    desc_run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    desc.space_after = Pt(25)
    
    # 按类别分组
    categories = {}
    for word in words:
        word_info = get_word_info(word)
        if word_info:
            cat = word_info["category"]
            if cat not in categories:
                categories[cat] = []
            categories[cat].append(word)
    
    # 为每个类别生成内容
    category_names = {
        "animals": "🐾 动物篇",
        "food": "🍎 食物篇"
    }
    
    for category, category_words in categories.items():
        # 类别标题
        cat_title = doc.add_paragraph()
        cat_run = cat_title.add_run(category_names.get(category, category.title()))
        cat_run.font.size = Pt(26)
        cat_run.font.bold = True
        cat_run.font.color.rgb = RGBColor(0xE7, 0x4C, 0x3C)
        cat_title.space_before = Pt(15)
        cat_title.space_after = Pt(10)
        
        # 为每个单词生成短语和例句
        for word in category_words:
            word_info = get_word_info(word)
            
            # 单词标题
            word_title = doc.add_paragraph()
            word_title.space_before = Pt(12)
            
            # emoji和单词
            emoji_run = word_title.add_run(f"{word_info['emoji']} ")
            emoji_run.font.size = Pt(22)
            
            word_run = word_title.add_run(f"{word}")
            word_run.font.size = Pt(22)
            word_run.font.bold = True
            
            chinese_run = word_title.add_run(f" ({word_info['chinese']})")
            chinese_run.font.size = Pt(18)
            chinese_run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
            
            # 简单短语
            phrases_title = doc.add_paragraph()
            phrases_title.space_before = Pt(6)
            pt_run = phrases_title.add_run("   📝 简单短语：")
            pt_run.font.size = Pt(14)
            pt_run.font.bold = True
            pt_run.font.color.rgb = RGBColor(0x34, 0x98, 0xDB)
            
            for phrase in word_info["phrases"]:
                phrase_para = doc.add_paragraph()
                phrase_para.add_run("       • " + phrase)
                phrase_para.runs[0].font.size = Pt(13)
                phrase_para.space_before = Pt(3)
            
            # 实用例句
            sentences_title = doc.add_paragraph()
            sentences_title.space_before = Pt(6)
            st_run = sentences_title.add_run("   💬 实用例句：")
            st_run.font.size = Pt(14)
            st_run.font.bold = True
            st_run.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)
            
            for sentence in word_info["sentences"]:
                sentence_para = doc.add_paragraph()
                sentence_para.add_run("       • " + sentence)
                sentence_para.runs[0].font.size = Pt(13)
                sentence_para.space_before = Pt(3)
            
            # 分隔线
            doc.add_paragraph()
        
        # 分页
        if category != list(categories.keys())[-1]:
            doc.add_page_break()
    
    # 保存文档
    doc.save(output_filename)
    print(f"\n✅ 短语和例句文档已生成: {output_filename}")
    print(f"📄 包含 {len(words)} 个单词的短语和例句")
    return output_filename

if __name__ == "__main__":
    print("\n" + "="*60)
    print("🎨 英语单词卡片生成系统")
    print("="*60)
    
    # 生成单词卡片
    card_file = generate_word_cards()
    
    # 生成短语和例句文档
    phrases_file = generate_phrases_document()
    
    print("\n" + "="*60)
    print("✅ 所有文档生成完成！")
    print("="*60 + "\n")
